# utilities/doc_processor.py

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from typing import List, Dict, Tuple
import re
from datetime import datetime
import os
import json

class DocumentProcessor:
    """
    A utility class for processing .docx Word documents, handling template variables,
    and managing document formatting.
    """
    
    def __init__(self, anthropic_client=None):
        self.anthropic_client = anthropic_client
        
        # Known variable patterns to handle
        self.variable_patterns = {
            'customer_name': r'\[Customer\]',
            'customer_address': r'\[Customer Address\]',
            'effective_date': r'_{3,}|\[Effective Date\]',
            'contractor_name': r'_{3,}|\[Contractor\]',
            'contractor_address': r'_{3,}|\[Contractor Address\]'
        }

    def get_run_properties(self, run) -> Dict:
        """Capture all XML-level properties of a run for .docx"""
        properties = {}
        
        if run._element.rPr:
            rPr = run._element.rPr
            
            # Core formatting
            properties['bold'] = rPr.find(qn('w:b')) is not None
            properties['italic'] = rPr.find(qn('w:i')) is not None
            properties['underline'] = rPr.find(qn('w:u')) is not None
            
            # Font
            font_elem = rPr.find(qn('w:rFonts'))
            if font_elem is not None:
                properties['font'] = font_elem.get(qn('w:ascii'))
            
            # Size
            size_elem = rPr.find(qn('w:sz'))
            if size_elem is not None:
                properties['size'] = int(size_elem.get(qn('w:val'))) / 2
            
            # Color
            color_elem = rPr.find(qn('w:color'))
            if color_elem is not None:
                properties['color'] = color_elem.get(qn('w:val'))
        
        return properties

    def apply_run_properties(self, run, properties: Dict):
        """Apply XML-level properties to a run in .docx"""
        if not properties:
            return
            
        rPr = run._element.get_or_add_rPr()
        
        # Core formatting
        if properties.get('bold'):
            b = rPr.find(qn('w:b'))
            if b is None:
                b = OxmlElement('w:b')
                rPr.append(b)
        
        if properties.get('italic'):
            i = rPr.find(qn('w:i'))
            if i is None:
                i = OxmlElement('w:i')
                rPr.append(i)
                
        if properties.get('underline'):
            u = rPr.find(qn('w:u'))
            if u is None:
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
        
        # Font
        if 'font' in properties:
            fonts = rPr.find(qn('w:rFonts'))
            if fonts is None:
                fonts = OxmlElement('w:rFonts')
                rPr.append(fonts)
            fonts.set(qn('w:ascii'), properties['font'])
            fonts.set(qn('w:hAnsi'), properties['font'])
        
        # Size
        if 'size' in properties:
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz')
                rPr.append(sz)
            sz.set(qn('w:val'), str(int(properties['size'] * 2)))
        
        # Color
        if 'color' in properties:
            color = rPr.find(qn('w:color'))
            if color is None:
                color = OxmlElement('w:color')
                rPr.append(color)
            color.set(qn('w:val'), properties['color'])

    def process_template(self, doc_path: str, variables: Dict[str, str]) -> Document:
        """Process a .docx template document with given variables"""
        doc = Document(doc_path)
        
        # Process main document body
        for paragraph in doc.paragraphs:
            self._process_paragraph(paragraph, variables)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_paragraph(paragraph, variables)
        
        return doc

    def _process_paragraph(self, paragraph, variables: Dict[str, str]):
        """Process a single paragraph for variable replacement"""
        for var_name, pattern in self.variable_patterns.items():
            if var_name in variables:
                for run in paragraph.runs:
                    if re.search(pattern, run.text):
                        # Capture properties before replacement
                        properties = self.get_run_properties(run)
                        
                        # Replace text
                        run.text = re.sub(pattern, variables[var_name], run.text)
                        
                        # Reapply properties
                        self.apply_run_properties(run, properties)

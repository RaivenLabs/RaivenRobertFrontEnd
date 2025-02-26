# app/utils/document_processor.py
import os
import pandas as pd
import docx
import PyPDF2
import json
import re

def extract_text_from_file(file_path, file_extension):
    """
    Extract text content from various file types
    
    Args:
        file_path: Path to the file
        file_extension: Extension of the file (pdf, docx, txt)
        
    Returns:
        str: Extracted text content
    """
    if file_extension == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == 'docx':
        return extract_text_from_docx(file_path)
    elif file_extension == 'txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {file_extension}")

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
            
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n\n"
        
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        return f"Error: {str(e)}"

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Add text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(' | '.join(row_text))
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        return f"Error: {str(e)}"

def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting text from TXT with latin-1 encoding: {str(e)}")
            return f"Error: {str(e)}"
    except Exception as e:
        print(f"Error extracting text from TXT: {str(e)}")
        return f"Error: {str(e)}"

def parse_csv(file_path):
    """
    Parse CSV file and return pandas DataFrame
    
    Args:
        file_path: Path to the CSV file
        
    Returns:
        pandas.DataFrame: Parsed CSV data
    """
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                # First try with pandas auto-detection of separator
                df = pd.read_csv(file_path, encoding=encoding)
                return df
            except pd.errors.ParserError:
                # If that fails, try to detect the delimiter
                with open(file_path, 'r', encoding=encoding) as f:
                    sample = f.read(1024)
                    
                dialect = get_csv_dialect(sample)
                df = pd.read_csv(file_path, encoding=encoding, sep=dialect.delimiter)
                return df
            except UnicodeDecodeError:
                # Try next encoding
                continue
            except Exception as e:
                print(f"Error parsing CSV with encoding {encoding}: {str(e)}")
                continue
        
        # If all encodings fail, raise an exception
        raise Exception("Failed to parse CSV with any encoding")
    
    except Exception as e:
        print(f"Error parsing CSV: {str(e)}")
        raise

def get_csv_dialect(sample):
    """
    Detect CSV dialect (delimiter, quotechar, etc.)
    
    Args:
        sample: Sample of CSV content
        
    Returns:
        csv.Dialect: CSV dialect
    """
    import csv
    sniffer = csv.Sniffer()
    dialect = sniffer.sniff(sample)
    return dialect

def parse_excel(file_path):
    """
    Parse Excel file and return sheet data
    
    Args:
        file_path: Path to the Excel file
        
    Returns:
        dict: Sheet data with sheet names as keys
    """
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        sheet_names = excel_file.sheet_names
        
        sheets_data = {}
        for sheet_name in sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # Get column info and sample data
            columns = df.columns.tolist()
            sample_data = df.head(5).to_dict(orient='records')
            data_types = {col: str(df[col].dtype) for col in columns}
            
            sheets_data[sheet_name] = {
                "columns": columns,
                "data_types": data_types,
                "sample_data": sample_data,
                "row_count": len(df)
            }
        
        return sheets_data
    
    except Exception as e:
        print(f"Error parsing Excel: {str(e)}")
        raise

def extract_field_patterns(text):
    """
    Extract potential field patterns from text
    
    Args:
        text: Text content to analyze
        
    Returns:
        list: Potential field patterns found in the text
    """
    patterns = []
    
    # Look for "field: value" patterns
    field_value_pattern = re.compile(r'([A-Z][a-zA-Z\s]+):\s*(.+)(?:\n|$)')
    matches = field_value_pattern.findall(text)
    
    for field, value in matches:
        field_name = field.strip().lower().replace(' ', '_')
        patterns.append({
            "field_name": field_name,
            "example_value": value.strip(),
            "pattern": f"{field}: {value.strip()}"
        })
    
    # Look for form-like patterns
    form_pattern = re.compile(r'([A-Z][a-zA-Z\s]+)\s*[_-]{2,}\s*(.+)(?:\n|$)')
    matches = form_pattern.findall(text)
    
    for field, value in matches:
        field_name = field.strip().lower().replace(' ', '_')
        patterns.append({
            "field_name": field_name,
            "example_value": value.strip(),
            "pattern": f"{field} --- {value.strip()}"
        })
    
    return patterns

def detect_data_type(value):
    """
    Detect the likely data type of a value
    
    Args:
        value: Value to analyze
        
    Returns:
        str: Detected data type
    """
    if value is None:
        return "null"
    
    if isinstance(value, bool):
        return "boolean"
    
    if isinstance(value, (int, float)):
        return "number"
    
    if not isinstance(value, str):
        return "unknown"
    
    # Try to convert to date
    date_patterns = [
        r'\d{4}-\d{2}-\d{2}',  # YYYY-MM-DD
        r'\d{2}/\d{2}/\d{4}',  # MM/DD/YYYY
        r'\d{2}-\d{2}-\d{4}'   # MM-DD-YYYY
    ]
    
    for pattern in date_patterns:
        if re.fullmatch(pattern, value):
            return "date"
    
    # Check for boolean-like values
    bool_values = ['true', 'false', 'yes', 'no', 'y', 'n']
    if value.lower() in bool_values:
        return "boolean"
    
    # Check if it's a number
    try:
        float(value)
        return "number"
    except ValueError:
        pass
    
    # Check for UUID pattern
    uuid_pattern = r'[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}'
    if re.fullmatch(uuid_pattern, value, re.IGNORECASE):
        return "uuid"
    
    # Default to string
    return "string"
